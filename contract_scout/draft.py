"""Генерация проекта любого гражданско-правового договора (включая IT)."""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from datetime import date, datetime
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from contract_scout.llm import ChatLLM
from contract_scout.redact import public_brief
from contract_scout.types import kind_frame, kind_label, normalize_kind

DRAFT_SYSTEM = (
    "Ты составляешь статьи гражданско-правовых договоров по праву РФ для переговоров. "
    "Удачные примеры в запросе — системный ориентир по структуре, нумерации и защитным нормам. "
    "Адаптируй их под brief; не копируй реквизиты, ФИО и чужие суммы дословно. "
    "Без markdown (# * ** `), только нумерованные статьи."
)

DRAFT_PROMPT = """Составь ТЕКСТ СТАТЕЙ проекта договора по праву РФ. Тип: {kind_label} ({kind_frame}).
Это черновик для переговоров, не нотариальный акт и не замена юристу.

Формат строго:
- только статьи, начиная со строки «1. Предмет» или «## 1. Предмет»
- пункты: 1.1., 1.2.; подпункты: 1.1.1., 1.1.2.
- не пиши шапку «ДОГОВОР», город, дату, преамбулу сторон, реквизиты и подписи — их подставят локально
- стороны в тексте статей: «Заказчик»/«Исполнитель» или роли по типу договора, без названий организаций и ФИО
- без markdown: без # * ** ` и маркированных списков

Защитные нормы (если уместны типу):
- существенные условия названы явно;
- цена и НДС/налоговый режим понятны;
- зеркальная неустойка с потолком;
- лимит ответственности = цена договора, без упущенной выгоды (кроме умысла);
- одностороннее изменение условий запрещено;
- применимое право РФ, разумная подсудность.

Если тип IT / лицензия / разработка:
- IP: заказчику результат после оплаты; фоновые библиотеки у исполнителя + лицензия;
- 2 круга правок в рамках ТЗ; change request;
- автоприёмка через 7 рабочих дней молчания.
- если в brief include_tz / include_act = true — сошлись в статьях на приложения №…, сами тексты приложений не пиши.

Если купля-продажа / поставка: переход риска, скрытые недостатки, комплектность.
Если аренда: предмет, срок, плата, формула индексации, возврат.
Если NDA: состав тайны, срок 2–3 года, исключения, возврат носителей.
Если заём: ставка, срок возврата, один вид санкции за просрочку.
Если ГПХ: оплата за результат, без подчинения ПВТР.

Данные (без реквизитов сторон):
{brief}

Удачные примеры из библиотеки (обезличенные статьи). Используй как системный ориентир по структуре,
нумерации и защитным нормам. Не копируй дословно чужие цены/сроки — подставь данные из brief.
Не переноси реквизиты, ФИО и названия организаций из примеров.
{examples}

Верни только статьи. Заголовки статей пиши как: 1. Предмет (без решёток #).
Пункты: 1.1., 1.1.1. Без *, **, `, списков с дефисами и без блока реквизитов.
"""

REVISE_PROMPT = """Ты правишь СТАТЬИ проекта договора по праву РФ. Это черновик для переговоров, не замена юристу.

Инструкция пользователя (что добавить / убрать / изменить):
{instruction}

Правила:
- верни ТОЛЬКО статьи договора, начиная с «1. …» или «## 1. …»;
- сохрани стиль нумерации: 1. / 1.1. / 1.1.1.;
- если добавляешь статью — вставь с новой нумерацией и перенумеруй последующие;
- не пиши шапку «ДОГОВОР», город, дату, преамбулу сторон, реквизиты и подписи;
- стороны называй ролями («Заказчик», «Исполнитель» и т.п.), без ИНН, счетов и ФИО;
- без markdown-мусора: без # * ** ` и маркированных списков;
- не выкидывай защитные нормы без явной просьбы пользователя.

Текущие статьи:
{articles}
"""

FIX_RISKS_PROMPT = """Ты правишь текст гражданско-правового договора по праву РФ.
Это черновик для переговоров, не нотариальный акт и не замена юристу.

Задача: устранить или смягчить найденные УЗКИЕ МЕСТА по рекомендациям «как чинить»
и при необходимости добавить недостающие условия.

Найденные риски:
{risks}

Чего может не хватать:
{missing}

Правила:
- верни ПОЛНЫЙ исправленный текст договора (шапка/преамбула/статьи/заключения — если они были);
- сохрани смысл сделки, сроки и суммы из исходника, если риск их не требует менять;
- стороны — ролями («Заказчик», «Исполнитель» и т.п.); плейсхолдеры вроде [ОРГАНИЗАЦИЯ], [ИНН], [СЧЁТ] не восстанавливай;
- нумерация: 1. / 1.1. / 1.1.1.;
- без markdown: без # * ** ` и маркированных списков с дефисами;
- не выдумывай новые реквизиты, ФИО и банковские счета.

Исходный текст:
{contract}
"""


@dataclass
class DraftBrief:
    contract_kind: str = "it"
    customer_name: str = ""
    customer_person_type: str = "ooo"
    customer_inn_kpp: str = ""
    customer_address: str = ""
    customer_phone: str = ""
    customer_email: str = ""
    customer_rs: str = ""
    customer_bank: str = ""
    customer_bik: str = ""
    customer_ks: str = ""
    customer_form_label: str = ""
    contractor_name: str = ""
    contractor_person_type: str = "ooo"
    contractor_inn_kpp: str = ""
    contractor_address: str = ""
    contractor_phone: str = ""
    contractor_email: str = ""
    contractor_rs: str = ""
    contractor_bank: str = ""
    contractor_bik: str = ""
    contractor_ks: str = ""
    contractor_form_label: str = ""
    subject: str = "согласованные сторонами услуги / работы / товар"
    scope: str = "согласно приложению № 1"
    price: str = "500 000"
    currency: str = "руб."
    prepay_percent: str = "40"
    term_days: str = "45"
    city: str = "Москва"
    extra: str = ""
    extra_fields: Dict[str, str] = field(default_factory=dict)
    contract_number: str = "б/н"
    contract_date: str = ""
    customer_ogrn: str = ""
    customer_rep_title: str = ""
    customer_rep: str = ""
    customer_basis: str = ""
    contractor_ogrn: str = ""
    contractor_rep_title: str = ""
    contractor_rep: str = ""
    contractor_basis: str = ""
    # старые поля формы — складываем в inn/address, если новые пусты
    customer_details: str = ""
    contractor_details: str = ""
    include_tz: bool = False
    include_act: bool = False


def _as_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"1", "true", "yes", "on"}


def brief_from_form(data: Dict[str, Any]) -> DraftBrief:
    fields = {k: ("" if v is None else str(v)).strip() for k, v in data.items()}
    known = {f.name for f in DraftBrief.__dataclass_fields__.values() if f.name != "extra_fields"}
    keep_default_if_empty = {
        "price",
        "prepay_percent",
        "term_days",
        "city",
        "subject",
        "scope",
        "currency",
        "contract_kind",
        "customer_person_type",
        "contractor_person_type",
        "contract_number",
    }
    bool_keys = {"include_tz", "include_act"}
    kwargs = {}
    for key, value in fields.items():
        if key not in known or key in bool_keys:
            continue
        if not value and key in keep_default_if_empty:
            continue
        kwargs[key] = value
    kwargs["include_tz"] = _as_bool(data.get("include_tz"))
    kwargs["include_act"] = _as_bool(data.get("include_act"))
    brief = DraftBrief(**kwargs)
    if brief.customer_details and not brief.customer_inn_kpp:
        brief.customer_inn_kpp = brief.customer_details
    if brief.contractor_details and not brief.contractor_inn_kpp:
        brief.contractor_inn_kpp = brief.contractor_details
    kind = normalize_kind(brief.contract_kind, f"{brief.subject} {brief.scope} {brief.extra}")
    if kind == "any" and (brief.contract_kind or "auto") in {"auto", ""}:
        kind = "services"
    brief.contract_kind = kind
    return brief


def _roles(kind: str) -> tuple[str, str]:
    return {
        "sale": ("Продавец", "Покупатель"),
        "supply": ("Поставщик", "Покупатель"),
        "lease": ("Арендодатель", "Арендатор"),
        "loan": ("Займодавец", "Заёмщик"),
        "nda": ("Раскрывающая сторона", "Получающая сторона"),
        "agency": ("Принципал", "Агент"),
        "license": ("Лицензиар", "Лицензиат"),
    }.get(kind, ("Заказчик", "Исполнитель"))


def fallback_markdown(brief: DraftBrief) -> str:
    kind = normalize_kind(brief.contract_kind, brief.subject)
    if kind == "it":
        body, last = _it_body(brief)
        title = "ДОГОВОР возмездного оказания услуг (IT)"
    elif kind == "nda":
        body, last = _nda_body(brief)
        title = "СОГЛАШЕНИЕ о конфиденциальности"
    elif kind == "sale" or kind == "supply":
        body, last = _sale_body(brief, kind)
        title = "ДОГОВОР поставки" if kind == "supply" else "ДОГОВОР купли-продажи"
    elif kind == "lease":
        body, last = _lease_body(brief)
        title = "ДОГОВОР аренды"
    elif kind == "loan":
        body, last = _loan_body(brief)
        title = "ДОГОВОР займа"
    else:
        body, last = _generic_body(brief, kind)
        title = f"ДОГОВОР ({kind_label(kind)})"
    return _assemble_draft(brief, title, body, last)


def _assemble_draft(brief: DraftBrief, title: str, body: str, last_article: int) -> str:
    closing, req_n = _closing(brief, last_article + 1)
    apps = _appendices_markdown(brief)
    return (
        _header(brief, title)
        + body
        + closing
        + apps
        + f"\n## {req_n}. Реквизиты и подписи сторон\n"
    )


def _party_name(value: str, placeholder: str) -> str:
    text = (value or "").strip()
    return text or placeholder


def _format_party_name(person_type: str, name: str, form_label: str = "") -> str:
    """Для ООО в поле — только название («Вектор»), в тексте договора — ООО «Вектор»."""
    text = (name or "").strip()
    if not text:
        return "________________"
    key = (person_type or "ooo").lower().strip()
    up = text.upper().replace("Ё", "Е")
    if key == "ooo":
        if up.startswith("ООО") or "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ" in up:
            return text
        core = text.strip("«»\"' ")
        return f"ООО «{core}»"
    if key == "custom" and (form_label or "").strip():
        label = (form_label or "").strip()
        if up.startswith(label.upper().replace("Ё", "Е")):
            return text
        return text
    return text


PARTY_FORMS = {
    "ooo": {
        "title": "ООО",
        "name_label": "Наименование",
        "inn_label": "ИНН/КПП",
        "addr_label": "Юр. адрес",
        "legal": True,
        "name_ph": "Вектор",
        "inn_ph": "7700000000 / 770001001",
    },
    "legal": {
        "title": "ООО",
        "name_label": "Наименование",
        "inn_label": "ИНН/КПП",
        "addr_label": "Юр. адрес",
        "legal": True,
        "name_ph": "Вектор",
        "inn_ph": "7700000000 / 770001001",
    },
    "ip": {
        "title": "ИП",
        "name_label": "ФИО / наименование ИП",
        "inn_label": "ИНН",
        "addr_label": "Адрес",
        "legal": False,
        "name_ph": "ИП Иванов Иван Иванович",
        "inn_ph": "770000000000",
    },
    "selfemployed": {
        "title": "самозанятый",
        "name_label": "ФИО",
        "inn_label": "ИНН",
        "addr_label": "Адрес",
        "legal": False,
        "name_ph": "Иванов Иван Иванович",
        "inn_ph": "770000000000",
    },
    "individual": {
        "title": "физическое лицо",
        "name_label": "ФИО",
        "inn_label": "ИНН",
        "addr_label": "Адрес",
        "legal": False,
        "name_ph": "Иванов Иван Иванович",
        "inn_ph": "770000000000",
    },
    "custom": {
        "title": "иная форма",
        "name_label": "Наименование / ФИО",
        "inn_label": "ИНН / ИНН/КПП",
        "addr_label": "Адрес",
        "legal": True,
        "name_ph": "как в уставе / паспорте",
        "inn_ph": "",
    },
}


def _party_form(person_type: str) -> dict:
    key = (person_type or "ooo").lower().strip()
    return PARTY_FORMS.get(key, PARTY_FORMS["ooo"])


def _is_legal(person_type: str) -> bool:
    return bool(_party_form(person_type)["legal"])


def _brief_date(brief: DraftBrief) -> date:
    raw = (brief.contract_date or "").strip()
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(raw, fmt).date()
        except ValueError:
            continue
    return date.today()


def _ru_date(day: date) -> str:
    months = (
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    )
    return f"{day.day} {months[day.month - 1]} {day.year} г."


def _strip_md_noise(text: str) -> str:
    """Убрать markdown-мусор из текста договора для DOCX."""
    out = text or ""
    out = out.replace("\u00a0", " ")
    out = re.sub(r"```+\w*", "", out)
    out = re.sub(r"`([^`]+)`", r"\1", out)
    out = re.sub(r"\*\*([^*]+)\*\*", r"\1", out)
    out = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", out)
    out = re.sub(r"__([^_\n]+)__", r"\1", out)
    # одиночные _markdown_ только если буквы внутри, не линия подписи
    out = re.sub(r"(?<![A-Za-zА-Яа-яЁё0-9_])_([A-Za-zА-Яа-яЁё][^_\n]{0,80}?)_(?![A-Za-zА-Яа-яЁё0-9_])", r"\1", out)
    out = re.sub(r"^#{1,6}\s*", "", out)
    out = re.sub(r"^>\s*", "", out)
    out = re.sub(r"^[-*+]\s+", "", out)
    out = out.replace("**", "").replace("##", "")
    out = re.sub(r"(?<![A-Za-zА-Яа-яЁё0-9])#(?!\d)", "", out)
    out = re.sub(r"(?<!\*)\*(?!\*)", "", out)
    out = re.sub(r'"([^"]{1,80})"', r"\1", out)
    out = re.sub(r"(?<![A-Za-zА-Яа-яЁё])'([^']{1,80})'(?![A-Za-zА-Яа-яЁё])", r"\1", out)
    out = re.sub(r"[ \t]{2,}", " ", out)
    return out.strip()


def _default_basis(person_type: str) -> str:
    key = (person_type or "ooo").lower().strip()
    if key in {"ooo", "legal", "custom"}:
        return "Устава"
    if key == "ip":
        return "листа записи ЕГРИП"
    return "паспорта гражданина РФ"


def _default_rep_title(person_type: str) -> str:
    key = (person_type or "ooo").lower().strip()
    if key in {"ooo", "legal", "custom"}:
        return "Генерального директора"
    return ""


def _ogrn_label(person_type: str) -> str:
    key = (person_type or "ooo").lower().strip()
    if key == "ip":
        return "ОГРНИП"
    if key in {"selfemployed", "individual"}:
        return ""
    return "ОГРН"


def _signatory(brief: DraftBrief, side: str) -> Dict[str, str]:
    prefix = "customer_" if side == "customer" else "contractor_"
    person_type = str(getattr(brief, f"{prefix}person_type") or "ooo")
    title = (getattr(brief, f"{prefix}rep_title") or "").strip() or _default_rep_title(person_type)
    rep = (getattr(brief, f"{prefix}rep") or "").strip()
    basis = (getattr(brief, f"{prefix}basis") or "").strip() or _default_basis(person_type)
    form_label = str(getattr(brief, f"{prefix}form_label") or "")
    name = _format_party_name(
        person_type,
        str(getattr(brief, f"{prefix}name") or ""),
        form_label,
    )
    return {
        "person_type": person_type,
        "name": name,
        "title": title,
        "rep": rep or "________________",
        "basis": basis,
        "form_label": form_label,
        "inn_kpp": str(getattr(brief, f"{prefix}inn_kpp") or ""),
        "ogrn": str(getattr(brief, f"{prefix}ogrn") or ""),
        "address": str(getattr(brief, f"{prefix}address") or ""),
        "phone": str(getattr(brief, f"{prefix}phone") or ""),
        "email": str(getattr(brief, f"{prefix}email") or ""),
        "rs": str(getattr(brief, f"{prefix}rs") or ""),
        "bank": str(getattr(brief, f"{prefix}bank") or ""),
        "bik": str(getattr(brief, f"{prefix}bik") or ""),
        "ks": str(getattr(brief, f"{prefix}ks") or ""),
    }


def _acting_phrase(sign: Dict[str, str], role: str) -> str:
    named = "именуемое" if _is_legal(sign["person_type"]) else "именуемый(ая)"
    if _is_legal(sign["person_type"]) or (sign["title"] and sign["rep"] and sign["rep"] != "________________"):
        title = sign["title"] or _default_rep_title("ooo")
        return (
            f"{sign['name']}, {named} в дальнейшем «{role}», "
            f"в лице {title} {sign['rep']}, действующего на основании {sign['basis']}"
        )
    return (
        f"{sign['name']}, {named} в дальнейшем «{role}», "
        f"действующий(ая) от собственного имени на основании {sign['basis']}"
    )


def _party_block(
    title: str,
    *,
    person_type: str,
    name: str,
    inn_kpp: str,
    address: str,
    phone: str,
    email: str,
    rs: str,
    bank: str,
    bik: str,
    ks: str,
    form_label: str = "",
    ogrn: str = "",
    rep_title: str = "",
    rep: str = "",
    basis: str = "",
) -> str:
    form = _party_form(person_type)
    form_title = (form_label or "").strip() or form["title"]
    sign = {
        "person_type": person_type,
        "name": _party_name(name, "________________"),
        "title": (rep_title or "").strip() or _default_rep_title(person_type),
        "rep": (rep or "").strip() or "________________",
        "basis": (basis or "").strip() or _default_basis(person_type),
    }
    ogrn_l = _ogrn_label(person_type)
    ogrn_line = f"{ogrn_l}: {ogrn}\n" if ogrn_l else ""
    face = (
        f"В лице: {sign['title']} {sign['rep']}\nНа основании: {sign['basis']}"
        if _is_legal(person_type) or (rep_title or rep)
        else f"Действует лично на основании: {sign['basis']}"
    )
    return f"""### {title}
Форма: {form_title}
{form["name_label"]}: {name}
{face}
{form["inn_label"]}: {inn_kpp}
{ogrn_line}{form["addr_label"]}: {address}
Тел.: {phone}
Email: {email}
Банковские реквизиты:
Р/с: {rs}
Банк: {bank}
БИК: {bik}
К/с: {ks}
"""


def _header(brief: DraftBrief, title: str) -> str:
    when = _ru_date(_brief_date(brief))
    number = (brief.contract_number or "").strip() or "б/н"
    a, b = _roles(brief.contract_kind)
    left = _acting_phrase(_signatory(brief, "customer"), a)
    right = _acting_phrase(_signatory(brief, "contractor"), b)
    city = (brief.city or "").strip() or "________________"
    return f"""# {title} № {number}

г. {city} · {when}

{left}, с одной стороны, и {right}, с другой стороны, совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем.
"""


def _tail(brief: DraftBrief, article: int) -> str:
    """Обратная совместимость: заключительные + реквизиты (без приложений)."""
    closing, req = _closing(brief, article)
    return closing + f"\n## {req}. Реквизиты и подписи сторон\n"


def _annex_items(brief: DraftBrief) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    if brief.include_tz:
        items.append(("tz", "Техническое задание"))
    if brief.include_act:
        items.append(("act", "Акт выполненных работ (оказанных услуг)"))
    return items


def _closing(brief: DraftBrief, article: int) -> tuple[str, int]:
    extra = ""
    point = 4
    if (brief.extra or "").strip():
        extra = f"{article}.{point}. Особые условия: {brief.extra.strip()}\n"
        point += 1
    annex = _annex_items(brief)
    apps = ""
    if annex:
        listed = "; ".join(f"приложение № {i} — {title}" for i, (_, title) in enumerate(annex, 1))
        apps = (
            f"{article}.{point}. Неотъемлемыми частями настоящего Договора являются: {listed}.\n"
        )
        point += 1
    disputes = f"{article}.{point}. Споры рассматриваются в суде по месту нахождения ответчика, если императивные нормы не требуют иного."
    req = article + 1
    closing = f"""
## {article}. Заключительные положения
{article}.1. Настоящий Договор вступает в силу с даты его подписания Сторонами.
{article}.2. Применимое право — законодательство Российской Федерации.
{article}.3. Изменение условий — только письменным соглашением Сторон.
{extra}{apps}{disputes}
"""
    return closing, req


def _appendices_markdown(brief: DraftBrief) -> str:
    items = _annex_items(brief)
    if not items:
        return ""
    parts = []
    for num, (key, title) in enumerate(items, 1):
        if key == "tz":
            parts.append(_tz_appendix(brief, num, title))
        else:
            parts.append(_act_appendix(brief, num, title))
    return "\n" + "\n\n".join(parts) + "\n"


def _tz_appendix(brief: DraftBrief, num: int, title: str) -> str:
    number = (brief.contract_number or "").strip() or "б/н"
    when = _ru_date(_brief_date(brief))
    return f"""## ПРИЛОЖЕНИЕ № {num}
к Договору № {number} от {when}

{title}

1. Цель работ / услуг: {brief.subject}.
2. Состав и объём: {brief.scope}.
3. Срок выполнения: {brief.term_days} календарных дней.
4. Результат и критерии приёмки: ________________
5. Исходные данные, доступы и материалы Заказчика: ________________
6. Ограничения и допущения: ________________
7. Иные условия: {(brief.extra or "").strip() or "________________"}

Подписи сторон:

Заказчик: ________________ / ________________

Исполнитель: ________________ / ________________
"""


def _act_appendix(brief: DraftBrief, num: int, title: str) -> str:
    number = (brief.contract_number or "").strip() or "б/н"
    when = _ru_date(_brief_date(brief))
    city = (brief.city or "").strip() or "________________"
    customer = _signatory(brief, "customer")["name"]
    contractor = _signatory(brief, "contractor")["name"]
    return f"""## ПРИЛОЖЕНИЕ № {num}
к Договору № {number} от {when}

{title}

г. {city} · {when}

Заказчик: {customer}

Исполнитель: {contractor}

1. По Договору № {number} от {when} Исполнитель выполнил (оказал): {brief.subject}.
2. Объём: {brief.scope}.
3. Стоимость составляет {brief.price} ({brief.currency}); порядок оплаты — согласно Договору (аванс {brief.prepay_percent}%).
4. Работы (услуги) выполнены в полном объёме / с замечаниями: ________________
5. Претензий по объёму, качеству и срокам Заказчик не имеет / имеет: ________________

Настоящий Акт составлен в двух экземплярах, по одному для каждой Стороны.

Заказчик: ________________ / ________________

Исполнитель: ________________ / ________________
"""


def _generic_body(brief: DraftBrief, kind: str) -> tuple[str, int]:
    remain = 100 - int(brief.prepay_percent or 40)
    tz_line = ""
    if brief.include_tz:
        tz_line = "\n1.1.3. Детализация объёма — в Техническом задании (приложение № 1)."
    act_bit = "после подписания акта или автоприёмки (молчание 7 рабочих дней)"
    if brief.include_act:
        act_n = 2 if brief.include_tz else 1
        act_bit = (
            f"после подписания Акта (приложение № {act_n}) "
            "или автоприёмки (молчание 7 рабочих дней)"
        )
    body = f"""
## 1. Предмет
1.1. Стороны обязуются: {brief.subject}.
1.1.1. Объём: {brief.scope}.
1.1.2. Рамка отношений: {kind_frame(kind)}.{tz_line}
1.2. Существенные условия, не урегулированные настоящим Договором, согласовываются в приложениях, являющихся его неотъемлемой частью.

## 2. Срок
2.1. Срок исполнения — {brief.term_days} календарных дней, если иное не согласовано в приложении.
2.1.1. Просрочка кредитора продлевает срок на соответствующее число дней.

## 3. Цена и расчёты
3.1. Цена составляет {brief.price} ({brief.currency}).
3.1.1. Налоговый режим Стороны указывают в реквизитах; НДС выделяется отдельно, если применим.
3.2. Аванс {brief.prepay_percent}%, остаток {remain}% — {act_bit}.
3.3. Аванс засчитывается в оплату. При отказе контрагента от Договора неисполненная часть аванса возвращается.

## 4. Приёмка
4.1. Мотивированные замечания направляются в течение 7 рабочих дней со ссылкой на пункт Договора или приложения.
4.1.1. Молчание в указанный срок означает приёмку.

## 5. Ответственность
5.1. Совокупная ответственность каждой Стороны ограничивается ценой Договора, кроме умысла.
5.2. Упущенная выгода не возмещается.
5.3. Неустойка за просрочку — 0,1% просроченной суммы в день, не более 10% такой суммы, зеркально для обеих Сторон.
"""
    return body, 5


def _it_body(brief: DraftBrief) -> tuple[str, int]:
    remain = 100 - int(brief.prepay_percent or 40)
    tz_line = ""
    if brief.include_tz:
        tz_n = 1
        tz_line = (
            f"\n1.1.2. Детализация — в Техническом задании (приложение № {tz_n}), "
            "являющемся неотъемлемой частью Договора."
        )
    act_line = "после акта или автоприёмки"
    if brief.include_act:
        act_n = 2 if brief.include_tz else 1
        act_line = f"после подписания Акта (приложение № {act_n}) или автоприёмки"
    body = f"""
## 1. Предмет
1.1. Исполнитель обязуется оказать, а Заказчик принять и оплатить услуги: {brief.subject}.
1.1.1. Объём: {brief.scope}.{tz_line}

## 2. Срок
2.1. Срок оказания услуг — {brief.term_days} календарных дней с даты поступления аванса.
2.1.1. Просрочка Заказчика (материалы, доступы, приёмка) продлевает срок на соответствующее число дней.

## 3. Цена и расчёты
3.1. Цена Договора составляет {brief.price} ({brief.currency}).
3.2. Аванс {brief.prepay_percent}%, остаток {remain}% {act_line}.
3.3. Изменение объёма оформляется дополнительным соглашением (change request).

## 4. Приёмка
4.1. Срок проверки результата — 7 рабочих дней с даты передачи.
4.1.1. Молчание Заказчика означает приёмку.
4.1.2. В рамках согласованного ТЗ допускается не более двух кругов правок.

## 5. Интеллектуальные права
5.1. Исключительное право на результат переходит к Заказчику после полной оплаты соответствующего этапа.
5.2. Фоновые библиотеки, шаблоны и инструменты остаются у Исполнителя.
5.2.1. Заказчику предоставляется лицензия на их использование в составе результата.

## 6. Конфиденциальность
6.1. Срок охраны конфиденциальной информации — 3 года с даты раскрытия.
6.1.1. Персональные данные обрабатываются только при наличии поручения на обработку.

## 7. Ответственность
7.1. Ответственность Исполнителя ограничивается ценой настоящего Договора.
7.1.1. Упущенная выгода не возмещается, кроме случаев умысла.
7.2. Неустойка — 0,1% цены этапа в день, не более 10% цены этапа, зеркально для обеих Сторон.

## 8. Расторжение
8.1. При отказе Заказчика от Договора оплачивается фактически оказанное.
"""
    return body, 8


def _nda_body(brief: DraftBrief) -> tuple[str, int]:
    body = f"""
## 1. Предмет
1.1. Конфиденциальная информация: {brief.subject}.
1.1.1. {brief.scope}

## 2. Обязанности
2.1. Получающая сторона не раскрывает сведения третьим лицам и использует их только для согласованной цели.
2.2. Исключения:
2.2.1. сведения стали публичными не по вине получателя;
2.2.2. сведения разработаны самостоятельно;
2.2.3. раскрытие требуется законом.

## 3. Срок
3.1. Обязательства действуют {brief.term_days} месяцев с даты раскрытия соответствующей информации, но не менее 24 месяцев.

## 4. Ответственность
4.1. Возмещается доказанный прямой ущерб.
"""
    return body, 4


def _sale_body(brief: DraftBrief, kind: str) -> tuple[str, int]:
    body = f"""
## 1. Предмет
1.1. {brief.subject}.
1.1.1. Количество, ассортимент, качество: {brief.scope}.

## 2. Цена
2.1. Цена составляет {brief.price} ({brief.currency}).
2.1.1. НДС указывается отдельно при наличии.

## 3. Передача и риски
3.1. Срок передачи — {brief.term_days} дней.
3.1.1. Риск гибели переходит с момента передачи, если иное не согласовано.
3.2. Явные недостатки заявляются при приёмке.
3.2.1. Скрытые недостатки — в разумный срок, не менее 30 дней (или в срок, установленный законом).

## 4. Оплата
4.1. Аванс {brief.prepay_percent}%, остаток после передачи (отгрузки).
"""
    return body, 4


def _lease_body(brief: DraftBrief) -> tuple[str, int]:
    body = f"""
## 1. Предмет
1.1. Арендодатель передаёт: {brief.subject}.
1.1.1. Характеристики: {brief.scope}.

## 2. Срок и плата
2.1. Срок {brief.term_days} дней (или иной срок в приложении).
2.1.1. Плата {brief.price} ({brief.currency}) за период.
2.2. Индексация — только по согласованной формуле, не чаще одного раза в год, с уведомлением за 30 дней.

## 3. Возврат
3.1. Вещь возвращается в состоянии с учётом нормального износа по акту возврата.
"""
    return body, 3


def _loan_body(brief: DraftBrief) -> tuple[str, int]:
    body = f"""
## 1. Предмет
1.1. Займодавец передаёт {brief.price} ({brief.currency}).
1.1.1. Цель / условия: {brief.subject}. {brief.scope}

## 2. Возврат
2.1. Срок возврата — {brief.term_days} дней с даты передачи суммы.

## 3. Проценты и санкции
3.1. Процентная ставка указывается в особых условиях.
3.1.1. За просрочку применяется один вид санкции с разумным пределом, без сложения штрафа, пени и повышенных процентов за одно и то же нарушение.
"""
    return body, 3


def _extract_articles(text: str) -> str:
    cleaned = (text or "").replace("```markdown", "").replace("```", "")
    lines = cleaned.splitlines()
    start = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if re.match(r"^##\s*\d+", stripped) or re.match(r"^\d+\.\s+\S", stripped) or re.match(r"^1\.\s", stripped):
            start = i
            break
    chunk = "\n".join(lines[start:]).strip()
    chunk = re.split(r"^##\s*\d*\.?\s*Реквизит", chunk, maxsplit=1, flags=re.M | re.I)[0]
    chunk = re.split(r"^\d+\.\s*Реквизит", chunk, maxsplit=1, flags=re.M | re.I)[0]
    chunk = re.split(r"^##\s*\d*\.?\s*Подпис", chunk, maxsplit=1, flags=re.M | re.I)[0]
    chunk = re.split(r"^##\s*\d*\.?\s*Заключительн", chunk, maxsplit=1, flags=re.M | re.I)[0]
    return chunk.strip()


def _max_article(text: str) -> int:
    nums = [int(n) for n in re.findall(r"(?:^##\s*|^)(\d+)(?:\.\s|\.\s*[A-ЯA-ZА-я])", text or "", flags=re.M)]
    # заголовки статей: "## 3. Цена" или "3. Цена"
    article_heads = re.findall(r"^(?:##\s*)?(\d+)\.\s+[^\d]", text or "", flags=re.M)
    nums = [int(n) for n in article_heads]
    return max(nums) if nums else 0


def _contract_title(kind: str) -> str:
    return {
        "it": "ДОГОВОР возмездного оказания услуг (IT)",
        "nda": "СОГЛАШЕНИЕ о конфиденциальности",
        "sale": "ДОГОВОР купли-продажи",
        "supply": "ДОГОВОР поставки",
        "lease": "ДОГОВОР аренды",
        "loan": "ДОГОВОР займа",
    }.get(kind, f"ДОГОВОР ({kind_label(kind)})")


def _split_contract(markdown: str) -> tuple[str, str, str]:
    """Шапка | статьи | хвост (реквизиты)."""
    text = (markdown or "").replace("\r\n", "\n")
    lines = text.splitlines()
    start = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if re.match(r"^##\s*\d+", stripped) or re.match(r"^\d+\.\s+\S", stripped):
            start = i
            break
    body_and_tail = "\n".join(lines[start:])
    header = "\n".join(lines[:start]).rstrip()
    tail_match = re.search(
        r"(?m)^(#{1,3}\s*\d*\.?\s*)?(Реквизиты и подписи|Реквизиты сторон)\b",
        body_and_tail,
        flags=re.I,
    )
    if tail_match:
        body = body_and_tail[: tail_match.start()].rstrip()
        tail = body_and_tail[tail_match.start() :].rstrip()
    else:
        body = body_and_tail.rstrip()
        tail = ""
    return header, body, tail


class DraftPipeline:
    def __init__(self, llm: ChatLLM) -> None:
        self.llm = llm

    def generate(self, brief: DraftBrief, *, examples: str = "") -> str:
        local = fallback_markdown(brief)
        if not self.llm.settings.llm_enabled:
            return local
        examples_block = (examples or "").strip() or (
            "Примеров в библиотеке пока нет — опирайся только на правила выше и brief."
        )
        try:
            text = self.llm.complete(
                DRAFT_PROMPT.format(
                    kind_label=kind_label(brief.contract_kind),
                    kind_frame=kind_frame(brief.contract_kind),
                    brief=json.dumps(public_brief(asdict(brief)), ensure_ascii=False, indent=2),
                    examples=examples_block,
                ),
                temperature=0.3,
                max_tokens=8000,
                system=DRAFT_SYSTEM,
            )
            body = _extract_articles(text)
            last = _max_article(body)
            if last < 3 or "1.1" not in body:
                return local
            return _assemble_draft(brief, _contract_title(brief.contract_kind), "\n" + body, last)
        except Exception:
            return local

    def revise(self, markdown: str, instruction: str) -> str:
        instruction = (instruction or "").strip()
        if len(instruction) < 3:
            raise RuntimeError("Напишите, что нужно изменить в договоре.")
        if not self.llm.settings.llm_enabled:
            raise RuntimeError("ИИ выключен (LOCAL_ONLY или нет ключа). Правьте текст вручную.")
        header, body, tail = _split_contract(markdown)
        if len(body) < 40:
            raise RuntimeError("Сначала сгенерируйте или вставьте текст договора.")
        text = self.llm.complete(
            REVISE_PROMPT.format(instruction=instruction, articles=body),
            temperature=0.25,
            max_tokens=8000,
            redact_output=False,
        )
        revised = _extract_articles(text)
        if len(revised) < 40 or "1." not in revised:
            raise RuntimeError("ИИ не вернул статьи договора. Попробуйте переформулировать запрос.")
        parts = [p for p in (header, revised, tail) if p]
        return "\n\n".join(parts).strip() + "\n"

    def fix_risks(self, contract_text: str, report: Dict[str, Any]) -> str:
        if not self.llm.settings.llm_enabled:
            raise RuntimeError("ИИ выключен (нет ключа). Нельзя автоматически исправить договор.")
        text = (contract_text or "").strip()
        if len(text) < 80:
            raise RuntimeError("Слишком короткий текст договора для правки.")
        risks = _format_risks_for_fix(report)
        missing = _format_missing_for_fix(report)
        if not risks and not missing:
            raise RuntimeError("Нет найденных рисков для исправления. Сначала проверьте договор.")
        raw = self.llm.complete(
            FIX_RISKS_PROMPT.format(risks=risks or "—", missing=missing or "—", contract=text[:24000]),
            temperature=0.25,
            max_tokens=8000,
        )
        fixed = _strip_md_noise((raw or "").strip())
        if len(fixed) < 80:
            raise RuntimeError("ИИ не вернул исправленный текст. Попробуйте ещё раз.")
        return fixed + ("\n" if not fixed.endswith("\n") else "")


def _format_risks_for_fix(report: Dict[str, Any]) -> str:
    lines = []
    for i, b in enumerate(report.get("bottlenecks") or [], 1):
        if not isinstance(b, dict):
            continue
        title = str(b.get("title") or f"Риск {i}").strip()
        sev = str(b.get("severity") or "").strip()
        why = str(b.get("why") or "").strip()
        fix = str(b.get("fix") or "").strip()
        quote = str(b.get("quote") or "").strip()
        clause_ref = str(b.get("clause_ref") or "").strip()
        block = [f"{i}. {title}" + (f" [{sev}]" if sev else "")]
        if clause_ref:
            block.append(f"   Пункт договора: {clause_ref}")
        if quote:
            block.append(f"   Цитата: {quote}")
        if why:
            block.append(f"   Почему: {why}")
        if fix:
            block.append(f"   Как чинить: {fix}")
        lines.append("\n".join(block))
    return "\n".join(lines)


def _format_missing_for_fix(report: Dict[str, Any]) -> str:
    lines = []
    for i, m in enumerate(report.get("missing_clauses") or [], 1):
        if isinstance(m, dict):
            title = str(m.get("title") or "").strip()
            why = str(m.get("why") or "").strip()
            if title:
                lines.append(f"{i}. {title}" + (f" — {why}" if why else ""))
        elif m:
            lines.append(f"{i}. {m}")
    return "\n".join(lines)


def json_brief(brief: DraftBrief) -> str:
    return json.dumps(asdict(brief), ensure_ascii=False, indent=2)


def _set_run_font(run, *, bold: bool = False, size: int = 12) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def _clear_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")
        node.set(qn("w:sz"), "0")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "auto")


def _add_clause(doc: Document, text: str, *, indent: bool) -> None:
    clean = _strip_md_noise(text)
    if not clean:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
    _set_run_font(p.add_run(clean))


def _fill_sign_cell(cell, heading: str, lines: list[str]) -> None:
    cell.text = ""
    p0 = cell.paragraphs[0]
    r = p0.add_run(_strip_md_noise(heading))
    _set_run_font(r, bold=True, size=11)
    for line in lines:
        clean = _strip_md_noise(line)
        if not clean:
            cell.add_paragraph()
            continue
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _set_run_font(p.add_run(clean), size=11)


def _sign_lines(brief: DraftBrief, side: str) -> list[str]:
    sign = _signatory(brief, side)
    form = _party_form(sign["person_type"])
    form_title = (sign["form_label"] or "").strip() or form["title"]
    ogrn_l = _ogrn_label(sign["person_type"])
    signer = sign["rep"] if sign["rep"] != "________________" else sign["name"]
    lines = [
        f"Форма: {form_title}",
        f"{form['name_label']}: {sign['name']}",
    ]
    if _is_legal(sign["person_type"]) or sign["title"]:
        lines.append(f"В лице: {sign['title']} {sign['rep']}")
        lines.append(f"На основании: {sign['basis']}")
    else:
        lines.append(f"Действует лично на основании: {sign['basis']}")
    lines.append(f"{form['inn_label']}: {sign['inn_kpp'] or '________________'}")
    if ogrn_l:
        lines.append(f"{ogrn_l}: {sign['ogrn'] or '________________'}")
    lines.append(f"{form['addr_label']}: {sign['address'] or '________________'}")
    if sign["phone"]:
        lines.append(f"Тел.: {sign['phone']}")
    if sign["email"]:
        lines.append(f"Email: {sign['email']}")
    lines.append(f"Р/с: {sign['rs'] or '________________'}")
    lines.append(f"Банк: {sign['bank'] or '________________'}")
    lines.append(f"БИК: {sign['bik'] or '________________'}")
    lines.append(f"К/с: {sign['ks'] or '________________'}")
    lines.append("")
    lines.append(f"________________ / {signer}")
    lines.append("М.П.")
    return lines


def _add_city_date_row(doc: Document, city: str, when: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    _clear_table_borders(table)
    left, right = table.cell(0, 0), table.cell(0, 1)
    left.text = ""
    right.text = ""
    pl = left.paragraphs[0]
    pr = right.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(pl.add_run(f"г. {_strip_md_noise(city)}"))
    _set_run_font(pr.add_run(_strip_md_noise(when)))


def _add_requisites_table(doc: Document, brief: DraftBrief) -> None:
    a, b = _roles(brief.contract_kind)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    _set_run_font(p.add_run("Реквизиты и подписи сторон"), bold=True, size=12)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    _clear_table_borders(table)
    # тонкая вертикальная граница между колонками не нужна — стороны напротив
    _fill_sign_cell(table.cell(0, 0), a.upper(), _sign_lines(brief, "customer"))
    _fill_sign_cell(table.cell(0, 1), b.upper(), _sign_lines(brief, "contractor"))


def markdown_to_docx(markdown: str, dest, brief: DraftBrief | None = None) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    skip_rest = False
    for raw in (markdown or "").splitlines():
        stripped = raw.strip()
        if skip_rest:
            continue
        if stripped.startswith("CITY_DATE|"):
            parts = stripped.split("|", 2)
            city = parts[1] if len(parts) > 1 else ""
            when = parts[2] if len(parts) > 2 else ""
            _add_city_date_row(doc, city, when)
            continue
        city_date = re.match(r"^г\.\s*(.+?)\s+[·•|]\s*(.+)$", stripped)
        if city_date:
            _add_city_date_row(doc, city_date.group(1).strip(), city_date.group(2).strip())
            continue
        if re.match(r"^#{1,3}\s*\d*\.?\s*Реквизит", stripped, re.I) or "реквизиты и подписи" in stripped.lower():
            skip_rest = True
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title = _strip_md_noise(stripped[2:]).upper()
            _set_run_font(p.add_run(title), bold=True, size=14)
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            _set_run_font(p.add_run(_strip_md_noise(stripped[3:])), bold=True, size=12)
            continue
        if re.match(r"^\d+\.\s+[A-ЯA-Za-zА-яЁё]", stripped) and not re.match(r"^\d+\.\d+", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            _set_run_font(p.add_run(_strip_md_noise(stripped)), bold=True, size=12)
            continue
        if stripped.startswith("### "):
            continue
        if not stripped:
            continue
        if re.match(r"^\d+(?:\.\d+)+\.?\s", stripped):
            _add_clause(doc, stripped, indent=True)
            continue
        _add_clause(doc, stripped, indent=False)
    if brief is not None:
        _add_requisites_table(doc, brief)
    doc.save(dest)
