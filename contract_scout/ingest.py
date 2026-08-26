"""Локальный разбор PDF/DOCX/TXT без сети."""

from __future__ import annotations

import logging
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List
from xml.etree import ElementTree as ET

logger = logging.getLogger("contract_scout.ingest")

SUPPORTED = {".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".htm", ".rtf"}
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
EMPTY_MSG = (
    "Не удалось извлечь текст из файла. "
    "Сохраните договор как .docx или .txt. "
    "Скан PDF без текстового слоя сюда не подходит — нужен файл, где текст можно выделить мышью."
)


def chunk_text(text: str, size: int = 900, overlap: int = 120) -> List[str]:
    clean = re.sub(r"\s+\n", "\n", (text or "").replace("\r\n", "\n")).strip()
    if not clean:
        return []
    parts: List[str] = []
    i = 0
    while i < len(clean):
        piece = clean[i : i + size].strip()
        if piece:
            parts.append(piece)
        i += max(size - overlap, 1)
    return parts


def sniff_suffix(data: bytes, hint_name: str = "") -> str:
    """Определить расширение по magic bytes; иначе взять из имени файла."""
    raw = data[:16] if data else b""
    if raw.startswith(b"%PDF"):
        return ".pdf"
    if raw.startswith(b"PK"):
        # zip: docx/xlsx/odt — для договоров считаем docx (extract разберёт word/)
        return ".docx"
    if raw.startswith(b"{\\rtf") or raw.startswith(b"{\\rt"):
        return ".rtf"
    if raw[:4] == b"\xd0\xcf\x11\xe0":
        return ".doc"
    hint = Path(hint_name or "").suffix.lower()
    if hint in SUPPORTED:
        return hint
    # текстовые без magic — только по расширению или UTF-тексту
    if hint:
        return hint
    sample = data[:4000] if data else b""
    if sample and b"\x00" not in sample[:200]:
        try:
            sample.decode("utf-8")
            return ".txt"
        except UnicodeDecodeError:
            try:
                sample.decode("cp1251")
                return ".txt"
            except UnicodeDecodeError:
                pass
    return ""


def _sniff(path: Path) -> str:
    suf = sniff_suffix(path.read_bytes()[:4096], path.name)
    if suf == ".doc":
        return "ole"
    return suf.lstrip(".") or (path.suffix or "").lower().lstrip(".")


def _read_text_file(path: Path) -> str:
    data = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp1251", "cp866", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _docx_xml_text(path: Path) -> str:
    parts: List[str] = []
    with zipfile.ZipFile(path) as zf:
        names = [n for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml")]
        for name in names:
            if not any(key in name for key in ("document", "header", "footer", "footnotes", "comments")):
                continue
            root = ET.fromstring(zf.read(name))
            for node in root.iter(f"{W_NS}t"):
                if node.text:
                    parts.append(node.text)
                if node.tail:
                    parts.append(node.tail)
    return re.sub(r"[ \t]+", " ", "\n".join(parts))


def _docx_python(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    blocks: List[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                if p.text and p.text.strip():
                    blocks.append(p.text)
    return "\n".join(blocks)


def _extract_docx(path: Path) -> str:
    texts = []
    try:
        texts.append(_docx_python(path))
    except Exception as exc:  # noqa: BLE001
        logger.warning("python-docx failed: %s", exc)
    try:
        texts.append(_docx_xml_text(path))
    except Exception as exc:  # noqa: BLE001
        logger.warning("docx xml failed: %s", exc)
    return max(texts, key=lambda s: len(s or ""), default="") or ""


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: List[str] = []
    for page in reader.pages:
        got = ""
        try:
            got = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            got = ""
        if len(got.strip()) < 20:
            try:
                got = page.extract_text(extraction_mode="layout") or got
            except Exception:  # noqa: BLE001
                pass
        pages.append(got)
    return "\n".join(pages)


def _extract_rtf(path: Path) -> str:
    raw = _read_text_file(path)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    text = re.sub(r"\\[a-zA-Z]+\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", text).strip()


def _extract_ole_doc(path: Path) -> str:
    """Грубый разбор старого .doc: вытащить читаемые куски UTF-16/CP1251."""
    data = path.read_bytes()
    chunks: List[str] = []
    try:
        u16 = data.decode("utf-16le", errors="ignore")
        readable = "".join(ch if (ch.isprintable() or ch in "\n\r\t") else " " for ch in u16)
        readable = re.sub(r" {3,}", "\n", readable)
        chunks.append(readable)
    except Exception:  # noqa: BLE001
        pass
    chunks.append(_read_text_file(path))
    best = max(chunks, key=lambda s: len(re.findall(r"[А-Яа-яA-Za-z]{4,}", s or "")))
    letters = re.findall(r"[А-Яа-яA-Za-z]{3,}", best or "")
    return best if len(letters) >= 20 else ""


def extract_text(path: Path) -> str:
    kind = _sniff(path)
    logger.info("extract sniff=%s suffix=%s size=%s", kind, path.suffix, path.stat().st_size)
    if kind == "pdf" or path.suffix.lower() == ".pdf":
        return _extract_pdf(path)
    if kind == "docx":
        return _extract_docx(path)
    if kind == "rtf" or path.suffix.lower() == ".rtf":
        return _extract_rtf(path)
    if kind == "ole" or path.suffix.lower() == ".doc":
        if path.read_bytes()[:2] == b"PK":
            return _extract_docx(path)
        return _extract_ole_doc(path)
    if path.suffix.lower() == ".docx":
        return _extract_docx(path)
    return _read_text_file(path)


def ingest_path(path: Path, *, filename: str = "", local_only: bool = True) -> List[Dict[str, Any]]:
    path = Path(path)
    name = filename or path.name
    if not path.is_file() or path.stat().st_size == 0:
        raise RuntimeError("Файл пустой или не сохранился на диск.")
    text = extract_text(path)
    chunks = chunk_text(text)
    logger.info("ingest file=%s chars=%s chunks=%s local_only=%s", name, len(text or ""), len(chunks), local_only)
    if not chunks:
        raise RuntimeError(EMPTY_MSG)
    return [
        {
            "content": content,
            "metadata": {"filename": name, "chunk": idx, "kind": "contract"},
        }
        for idx, content in enumerate(chunks)
    ]
