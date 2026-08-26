from contract_scout.draft import DraftBrief, fallback_markdown
from contract_scout.knowledge import missing_clause_hints, scan_rules
from contract_scout.review import parse_json_object, rule_only_report, score_from_hits
from contract_scout.types import detect_kind


def test_detect_it_sample():
    assert detect_kind(PathText()) == "it"


def test_scan_finds_uncapped_liability():
    text = PathText()
    hits = scan_rules(text)
    ids = {h["id"] for h in hits}
    assert "liability_uncapped" in ids
    assert "ip_background" in ids
    assert "unlimited_revisions" in ids
    assert "personal_guarantee" in ids


def PathText() -> str:
    from pathlib import Path

    return (Path(__file__).resolve().parent.parent / "samples" / "risky_it_contract.txt").read_text(
        encoding="utf-8"
    )


def test_score_is_high_for_sample():
    hits = scan_rules(PathText())
    assert score_from_hits(hits) >= 70


def test_missing_clauses_on_empty():
    missing = missing_clause_hints("привет")
    assert len(missing) >= 8


def test_parse_json_from_fences():
    data = parse_json_object('```json\n{"overall_score": 12, "verdict": "низкий риск"}\n```')
    assert data["overall_score"] == 12


def test_rule_report_structure():
    report = rule_only_report(PathText(), "demo.txt")
    assert report["mode"] == "rules"
    assert report["bottlenecks"]
    assert "юридическ" in report["disclaimer"]


def test_fallback_sale_has_hidden_defects():
    md = fallback_markdown(DraftBrief(contract_kind="sale", subject="продажа станка"))
    assert "скрыт" in md.lower()


def test_fallback_contract_has_liability_cap():
    from contract_scout.draft import _sign_lines

    md = fallback_markdown(DraftBrief())
    assert "лимит" in md.lower() or "ответственност" in md.lower()
    lines = "\n".join(_sign_lines(DraftBrief(), "customer"))
    assert "ИНН/КПП" in lines
    assert "Р/с:" in lines
    assert "К/с:" in lines


def test_extract_docx_table(tmp_path):
    from docx import Document

    from contract_scout.ingest import extract_text, ingest_path, sniff_suffix

    path = tmp_path / "t.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Предмет договора"
    table.cell(0, 1).text = "разработка программного обеспечения"
    doc.save(path)
    text = extract_text(path)
    assert "Предмет" in text
    rows = ingest_path(path, filename="t.docx")
    assert rows
    assert sniff_suffix(path.read_bytes(), "без_расширения") == ".docx"
    assert sniff_suffix(b"%PDF-1.4", "x.bin") == ".pdf"


def test_party_block_stays_local():
    from dataclasses import asdict

    from contract_scout.draft import _sign_lines
    from contract_scout.redact import public_brief

    brief = DraftBrief(
        customer_name="ООО «Север»",
        customer_inn_kpp="7700000000 / 770001001",
        customer_rs="40702810100000000001",
        customer_bik="044525225",
    )
    md = fallback_markdown(brief)
    assert "7700000000" not in md
    assert "40702810100000000001" not in md
    assert "ООО «Север»" in md
    local = "\n".join(_sign_lines(brief, "customer"))
    assert "7700000000" in local
    assert "40702810100000000001" in local
    safe = public_brief(asdict(brief))
    assert "7700000000" not in str(safe)
    assert "customer_inn_kpp" not in safe


def test_individual_party_labels():
    from contract_scout.draft import _sign_lines

    cust = "\n".join(
        _sign_lines(DraftBrief(customer_person_type="individual", contractor_person_type="ooo"), "customer")
    )
    contr = "\n".join(
        _sign_lines(DraftBrief(customer_person_type="individual", contractor_person_type="ooo"), "contractor")
    )
    assert "Форма: физическое лицо" in cust
    assert "ФИО:" in cust
    assert "ИНН:" in cust
    assert "Адрес:" in cust
    assert "Форма: ООО" in contr
    assert "ИНН/КПП:" in contr
    assert "Юр. адрес:" in contr


def test_ip_and_selfemployed_labels():
    from contract_scout.draft import _sign_lines

    brief = DraftBrief(customer_person_type="ip", contractor_person_type="selfemployed")
    ip = "\n".join(_sign_lines(brief, "customer"))
    se = "\n".join(_sign_lines(brief, "contractor"))
    assert "Форма: ИП" in ip
    assert "ФИО / наименование ИП:" in ip
    assert "Форма: самозанятый" in se
    assert "ИНН/КПП:" not in se


def test_custom_form_label():
    from contract_scout.draft import _sign_lines

    lines = "\n".join(
        _sign_lines(
            DraftBrief(customer_person_type="custom", customer_form_label="АО", customer_name="АО Восток"),
            "customer",
        )
    )
    assert "Форма: АО" in lines
    assert "Наименование / ФИО:" in lines


def test_preamble_has_face_basis_city_date():
    md = fallback_markdown(
        DraftBrief(
            customer_name="ООО «Север»",
            customer_rep="Иванов Иван Иванович",
            customer_rep_title="Генерального директора",
            customer_basis="Устава",
            contractor_person_type="individual",
            contractor_name="Петров Пётр Петрович",
            city="Казань",
            contract_date="2026-08-25",
            contract_number="42",
        )
    )
    assert "в лице Генерального директора Иванов Иван Иванович" in md
    assert "действующего на основании Устава" in md
    assert "действующий(ая) от собственного имени на основании паспорта" in md
    assert "г. Казань · 25 августа 2026 г." in md
    assert "№ 42" in md
    assert "1.1.1." in md


def test_docx_has_signature_table(tmp_path):
    from docx import Document

    from contract_scout.draft import markdown_to_docx

    brief = DraftBrief(
        customer_name="ООО «Север»",
        customer_rep="Иванов И.И.",
        contractor_name="ИП Юг",
        contractor_person_type="ip",
        city="Казань",
        contract_date="2026-08-25",
    )
    md = fallback_markdown(brief)
    dest = tmp_path / "d.docx"
    markdown_to_docx(md, dest, brief=brief)
    doc = Document(dest)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "в лице" in text
    assert doc.tables
    assert len(doc.tables) >= 2  # город/дата + реквизиты
    assert "г. Казань" in doc.tables[0].cell(0, 0).text
    assert "25 августа 2026" in doc.tables[0].cell(0, 1).text
    left = doc.tables[-1].cell(0, 0).text
    right = doc.tables[-1].cell(0, 1).text
    assert "На основании" in left
    assert "ИНН" in left
    assert "Р/с" in left
    assert "ИП Юг" in right or "Форма: ИП" in right
    assert "#" not in text
    assert "**" not in text
    assert "###" not in text
    assert "________________" in left


def test_split_contract_keeps_header_and_tail():
    from contract_scout.draft import _split_contract

    md = """# ДОГОВОР № 1

г. Казань · 1 января 2026 г.

ООО «А», именуемое «Заказчик», и ООО «Б», именуемое «Исполнитель».

## 1. Предмет
1.1. Услуги.
## 2. Цена
2.1. 100.
## 3. Реквизиты и подписи сторон
"""
    header, body, tail = _split_contract(md)
    assert "ДОГОВОР" in header
    assert "1. Предмет" in body
    assert "Реквизиты" in tail
    assert "Реквизиты" not in body


def test_extract_contract_title_and_soft_search():
    from contract_scout.projects import extract_contract_title
    from contract_scout.service import ContractScout, EXAMPLE_MAX_RISK
    from contract_scout.config import load_settings

    text = "# ДОГОВОР возмездного оказания услуг (IT) № б/н\n\n1. Предмет\n"
    assert "ДОГОВОР" in extract_contract_title(text).upper()
    assert "услуг" in extract_contract_title(text).lower()

    s = ContractScout(load_settings())
    hits = s.search_projects("услуги")["projects"]
    assert hits
    assert any("услуг" in str(h.get("title") or "").lower() for h in hits)
    hits2 = s.search_projects("договор")["projects"]
    assert hits2
    assert ContractScout._is_good_example({"overall_score": 20, "verdict": "низкий риск"})
    assert not ContractScout._is_good_example({"overall_score": EXAMPLE_MAX_RISK + 10, "verdict": "высокий"})


def test_party_book_search(tmp_path):
    from contract_scout.parties import PartyBook

    book = PartyBook(tmp_path / "parties.json")
    book.upsert({"name": "ООО Север", "inn_kpp": "7700000000", "person_type": "ooo"})
    found = book.search("север")
    assert found and found[0]["inn_kpp"] == "7700000000"
    from contract_scout.redact import redact_requisites

    raw = (
        "ООО «Северные сервисы», ИНН 7700000000, ОГРН 1234567890123, "
        "р/с 40702810100000000001, адрес: г. Москва, ул. Тверская, 1, "
        "директор Иванов Иван Иванович, email boss@example.com"
    )
    safe, n = redact_requisites(raw)
    assert n >= 4
    assert "7700000000" not in safe
    assert "40702810100000000001" not in safe
    assert "boss@example.com" not in safe
    assert "Северные сервисы" not in safe
    assert "Иванов Иван" not in safe


def test_project_archive_roundtrip(tmp_path):
    from contract_scout.projects import ProjectArchive, public_summary

    archive = ProjectArchive(tmp_path / "projects")
    rec = archive.add(kind="contract", filename="a.txt", text="предмет договора", contract_kind="it")
    listed = archive.list()
    assert listed[0]["id"] == rec["id"]
    got = archive.get(rec["id"])
    assert got and "предмет" in got["text"]
    updated = archive.update(rec["id"], status="ai", report={"overall_score": 80, "verdict": "высокий риск"})
    assert updated["status"] == "ai"
    summary = public_summary(updated)
    assert summary["overall_score"] == 80
    assert "text" not in summary


def test_billing_credits_idempotent(tmp_path):
    from contract_scout.billing import BillingLedger
    from contract_scout.session import issue_cookie, parse_cookie

    book = BillingLedger(tmp_path / "billing.json")
    vid = "visitor-1"
    assert book.try_consume(vid, 1) is False
    assert book.apply_succeeded_payment(
        payment_id="pay-1",
        visitor_id=vid,
        credits=5,
        amount="490.00",
        raw_status="succeeded",
    )
    assert book.credits(vid) == 5
    assert (
        book.apply_succeeded_payment(
            payment_id="pay-1",
            visitor_id=vid,
            credits=5,
            amount="490.00",
            raw_status="succeeded",
        )
        is False
    )
    assert book.credits(vid) == 5
    assert book.try_consume(vid, 1) is True
    assert book.credits(vid) == 4
    token = issue_cookie("secret", vid)
    assert parse_cookie("secret", token) == vid
    assert parse_cookie("other", token) is None


def test_llm_without_key_raises():
    from pytest import raises

    from contract_scout.llm import ChatLLM

    class _S:
        local_only = True
        llm_enabled = False
        openai_api_key = ""
        openai_base_url = ""
        chat_model = "none"
        redact_requisites = True

    with raises(RuntimeError, match="ключа"):
        ChatLLM(_S()).complete("секретный договор")
